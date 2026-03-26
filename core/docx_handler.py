from docx import Document
import os

class DocxHandler:
    @staticmethod
    def translate_docx(file_path, translator, output_path=None):
        doc = Document(file_path)
        
        # 1. Collect all translatable objects
        targets = []
        for para in doc.paragraphs:
            if para.text.strip():
                targets.append(para)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.text.strip():
                            targets.append(para)
        
        # 2. Extract text and batch translate
        texts_to_translate = [t.text for t in targets]
        translated_texts = translator.translate_texts(texts_to_translate)
        
        # 3. Apply back
        for para, translated_text in zip(targets, translated_texts):
            if para.runs:
                for run in para.runs:
                    run.text = ""
                para.runs[0].text = translated_text
            else:
                para.add_run(translated_text)

        if output_path is None:
            output_path = file_path.replace(".docx", "_translated.docx")
        
        doc.save(output_path)
        return output_path
